# -*- coding: utf-8 -*-
"""تست‌های واحد برای تابع convert() در docx_to_json.py"""
import os
import sys
import tempfile

import pytest
from docx import Document

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
import docx_to_json  # noqa: E402


def make_sample_docx(path: str) -> None:
    doc = Document()
    doc.add_paragraph("اصفهان", style="Heading 1")
    doc.add_paragraph("عاشورا", style="Heading 2")
    doc.add_paragraph("شمر", style="Heading 3")
    doc.add_paragraph("ورود", style="Heading 4")
    doc.add_paragraph("بیت اول")
    doc.add_paragraph("بیت دوم")
    doc.add_paragraph("شهادت", style="Heading 4")
    doc.add_paragraph("بیت سوم")
    doc.save(path)


@pytest.fixture
def sample_docx_path(tmp_path):
    path = str(tmp_path / "sample.docx")
    make_sample_docx(path)
    return path


def test_convert_builds_correct_hierarchy(sample_docx_path):
    result = docx_to_json.convert(sample_docx_path)

    assert len(result) == 1
    field = result[0]
    assert field["title"] == "اصفهان"
    assert len(field["taziehs"]) == 1

    tazieh = field["taziehs"][0]
    assert tazieh["title"] == "عاشورا"
    assert len(tazieh["roles"]) == 1

    role = tazieh["roles"][0]
    assert role["title"] == "شمر"
    assert len(role["sections"]) == 2

    entry_section, martyrdom_section = role["sections"]
    assert entry_section["title"] == "ورود"
    assert entry_section["content"] == "بیت اول\nبیت دوم"
    assert martyrdom_section["title"] == "شهادت"
    assert martyrdom_section["content"] == "بیت سوم"


def test_convert_handles_missing_higher_headings(tmp_path):
    # فایلی که مستقیم با Heading 3 (نقش) شروع می‌شود، بدون زمینه/تعزیه صریح
    path = str(tmp_path / "no_field.docx")
    doc = Document()
    doc.add_paragraph("یزید", style="Heading 3")
    doc.add_paragraph("ورود", style="Heading 4")
    doc.add_paragraph("متن نمونه")
    doc.save(path)

    result = docx_to_json.convert(path)
    assert result[0]["title"] == "بدون زمینه"
    assert result[0]["taziehs"][0]["title"] == "بدون تعزیه"
    assert result[0]["taziehs"][0]["roles"][0]["title"] == "یزید"


def test_next_content_filename_increments_and_slugifies(tmp_path):
    content_dir = str(tmp_path / "content")
    os.makedirs(content_dir)
    # شبیه‌سازی دو فایل قبلی موجود
    open(os.path.join(content_dir, "001_ashura.json"), "w").close()
    open(os.path.join(content_dir, "002_bazar-e-sham.json"), "w").close()

    output_path = docx_to_json.next_content_filename(content_dir, "my new majles!.docx")

    assert os.path.dirname(output_path) == content_dir
    assert os.path.basename(output_path) == "003_my-new-majles.json"


def test_next_content_filename_starts_at_one_when_dir_empty(tmp_path):
    content_dir = str(tmp_path / "content")
    output_path = docx_to_json.next_content_filename(content_dir, "test.docx")
    assert os.path.basename(output_path) == "001_test.json"
